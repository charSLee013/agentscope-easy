# -*- coding: utf-8 -*-
# pylint: disable=W0212
"""The Word reader to read and chunk Word documents."""
import base64
import hashlib
import json
from typing import Literal, TYPE_CHECKING

from ._reader_base import ReaderBase
from ._text_reader import TextReader
from .._document import Document, DocMetadata
from ..._logging import logger
from ...message import ImageBlock, Base64Source, TextBlock

if TYPE_CHECKING:
    from docx.table import Table as DocxTable
    from docx.text.paragraph import Paragraph as DocxParagraph
else:
    DocxTable = "docx.table.Table"
    DocxParagraph = "docx.text.paragraph.Paragraph"

# VML (Vector Markup Language) namespace URI.
# Not registered in python-docx's default nsmap, so we define it here
# instead of relying on qn("v:...") which may fail in some versions.
_VML_NS = "{urn:schemas-microsoft-com:vml}"


def _extract_text_from_paragraph(para: DocxParagraph) -> str:
    """Extract text from a paragraph, including text in text boxes and
    shapes."""
    text = ""

    from docx.oxml.ns import qn

    for t_elem in para._element.findall(".//" + qn("w:t")):
        if t_elem.text:
            text += t_elem.text

    if not text:
        text = para.text.strip()

    if not text:
        txbx_contents = para._element.findall(".//" + qn("w:txbxContent"))
        for txbx in txbx_contents:
            for p_elem in txbx.findall(".//" + qn("w:p")):
                for t_elem in p_elem.findall(".//" + qn("w:t")):
                    if t_elem.text:
                        text += t_elem.text

        # Check for VML text boxes
        vml_textboxes = para._element.findall(".//" + _VML_NS + "textbox")
        for vml_tb in vml_textboxes:
            for p_elem in vml_tb.findall(".//" + qn("w:p")):
                for t_elem in p_elem.findall(".//" + qn("w:t")):
                    if t_elem.text:
                        text += t_elem.text

    return text.strip()


def _extract_table_data(table: DocxTable) -> list[list[str]]:
    """Extract table data, handling merged cells and preserving line breaks."""
    from docx.oxml.ns import qn

    table_data = []
    for tr in table._element.findall(qn("w:tr")):
        row_data = []

        tcs = tr.findall(qn("w:tc"))
        for tc in tcs:
            paragraphs = []
            for p_elem in tc.findall(qn("w:p")):
                texts = []
                for t_elem in p_elem.findall(".//" + qn("w:t")):
                    if t_elem.text:
                        texts.append(t_elem.text)

                para_text = "".join(texts)
                if para_text:
                    paragraphs.append(para_text)

            cell_text = "\n".join(paragraphs)
            row_data.append(cell_text)

        table_data.append(row_data)

    return table_data


def _extract_image_data(para: DocxParagraph) -> list[ImageBlock]:
    """Extract image data from a paragraph."""
    images: list[ImageBlock] = []

    from docx.oxml.ns import qn

    drawings = para._element.findall(".//" + qn("w:drawing"))

    for drawing in drawings:
        blips = drawing.findall(".//" + qn("a:blip"))

        for blip in blips:
            embed = blip.get(qn("r:embed"))

            if embed:
                try:
                    image_part = para.part.related_parts[embed]
                    image_data = image_part.blob
                    image_base64 = base64.b64encode(image_data).decode("utf-8")
                    content_type = image_part.content_type

                    images.append(
                        ImageBlock(
                            type="image",
                            source=Base64Source(
                                type="base64",
                                data=image_base64,
                                media_type=content_type,
                            ),
                        ),
                    )
                except Exception as e:
                    logger.error("Failed to extract image: %s", e)

    picts = para._element.findall(".//" + qn("w:pict"))

    for pict in picts:
        imagedatas = pict.findall(".//" + _VML_NS + "imagedata")

        for imagedata in imagedatas:
            rel_id = imagedata.get(qn("r:id"))

            if rel_id:
                try:
                    image_part = para.part.related_parts[rel_id]
                    image_data = image_part.blob
                    image_base64 = base64.b64encode(image_data).decode("utf-8")

                    images.append(
                        ImageBlock(
                            type="image",
                            source=Base64Source(
                                type="base64",
                                data=image_base64,
                                media_type=image_part.content_type,
                            ),
                        ),
                    )
                except Exception as e:
                    logger.error("Failed to extract image from pict: %s", e)
    return images


class WordReader(ReaderBase):
    """The reader that supports reading text, image, and table content from
    Word documents (.docx files), and chunking the text content into smaller
    pieces.

    .. note:: The table content is extracted in Markdown format.
    """

    def __init__(
        self,
        chunk_size: int = 512,
        split_by: Literal["char", "sentence", "paragraph"] = "sentence",
        include_image: bool = False,
        separate_table: bool = False,
        table_format: Literal["markdown", "json"] = "markdown",
    ) -> None:
        """Initialize the Word reader."""
        if chunk_size <= 0:
            raise ValueError(
                f"The chunk_size must be positive, got {chunk_size}",
            )

        if split_by not in ["char", "sentence", "paragraph"]:
            raise ValueError(
                "The split_by must be one of 'char', 'sentence' or "
                f"'paragraph', got {split_by}",
            )

        if table_format not in ["markdown", "json"]:
            raise ValueError(
                "The table_format must be one of 'markdown' or 'json', "
                f"got {table_format}",
            )

        self.chunk_size = chunk_size
        self.split_by = split_by
        self.include_image = include_image
        self.separate_table = separate_table
        self.table_format = table_format

        self._text_reader = TextReader(
            self.chunk_size,
            self.split_by,
        )

    async def __call__(self, word_path: str) -> list[Document]:
        """Read a Word document, split it into chunks, and return a list of
        Document objects."""

        blocks = self._get_data_blocks(word_path)

        doc_id = self.get_doc_id(word_path)
        documents: list[Document] = []
        for block in blocks:
            if block["type"] == "text":
                for _ in await self._text_reader(block["text"]):
                    documents.append(
                        Document(
                            # pylint: disable=E1120,E1125
                            metadata=DocMetadata(
                                content=_.metadata.content,
                                doc_id=doc_id,
                                chunk_id=0,
                                total_chunks=0,
                            ),
                            # pylint: enable=E1120,E1125
                        ),
                    )

            elif block["type"] == "image":
                documents.append(
                    Document(
                        # pylint: disable=E1120,E1125
                        metadata=DocMetadata(
                            content=block,
                            doc_id=doc_id,
                            chunk_id=0,
                            total_chunks=1,
                        ),
                        # pylint: enable=E1120,E1125
                    ),
                )

        total_chunks = len(documents)
        for idx, doc in enumerate(documents):
            doc.metadata.chunk_id = idx
            doc.metadata.total_chunks = total_chunks

        return documents

    def _get_data_blocks(self, word_path: str) -> list[TextBlock | ImageBlock]:
        """Return a list of data blocks extracted from the Word document."""
        try:
            from docx import Document as DocxDocument
            from docx.oxml import CT_P, CT_Tbl
            from docx.text.paragraph import Paragraph
            from docx.table import Table
            from docx.oxml.ns import qn
        except ImportError as e:
            raise ImportError(
                "Please install python-docx to use the Word reader. "
                "You can install it by `pip install python-docx`.",
            ) from e

        doc = DocxDocument(word_path)
        last_type = None

        blocks: list[TextBlock | ImageBlock] = []
        for element in doc.element.body:
            if isinstance(element, CT_P):
                para = Paragraph(element, doc)

                text = _extract_text_from_paragraph(para)

                if self.include_image:
                    has_drawing = bool(
                        para._element.findall(".//" + qn("w:drawing")),
                    )
                    has_pict = bool(
                        para._element.findall(".//" + qn("w:pict")),
                    )

                    if has_drawing or has_pict:
                        blocks.extend(_extract_image_data(para))
                        last_type = "image"

                if last_type == "text" or (
                    last_type == "table" and not self.separate_table
                ):
                    blocks[-1]["text"] += "\n" + text
                else:
                    blocks.append(
                        TextBlock(
                            type="text",
                            text=text,
                        ),
                    )

                last_type = "text"

            elif isinstance(element, CT_Tbl):
                table_data = _extract_table_data(Table(element, doc))

                if self.table_format == "markdown":
                    text = self._table_to_markdown(table_data)
                else:
                    text = self._table_to_json(table_data)

                if not self.separate_table and last_type in ["text", "table"]:
                    blocks[-1]["text"] += "\n" + text
                else:
                    blocks.append(
                        TextBlock(
                            type="text",
                            text=text,
                        ),
                    )

                last_type = "table"

        return blocks

    @staticmethod
    def _table_to_markdown(table_data: list[list[str]]) -> str:
        if not table_data:
            return ""

        num_cols = len(table_data[0])
        md_table = ""

        header_row = "| " + " | ".join(table_data[0]) + " |\n"
        md_table += header_row

        separator_row = "| " + " | ".join(["---"] * num_cols) + " |\n"
        md_table += separator_row

        for row in table_data[1:]:
            data_row = "| " + " | ".join(row) + " |\n"
            md_table += data_row

        return md_table

    @staticmethod
    def _table_to_json(table_data: list[list[str]]) -> str:
        json_strs = [
            "<system-info>A table loaded as a JSON array:</system-info>",
        ]

        for row in table_data:
            json_strs.append(
                json.dumps(row, ensure_ascii=False),
            )

        return "\n".join(json_strs)

    def get_doc_id(self, word_path: str) -> str:
        """Generate a document ID based on the Word file path."""
        return hashlib.md5(word_path.encode("utf-8")).hexdigest()
